from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "BC.docx"

FIGURE_LABELS = [
    "Hình 1. Kiến trúc client-server của MiniChat",
    "Hình 2. Mô hình giao tiếp client - server",
    "Hình 3. Luồng xử lý gửi và nhận tin nhắn trong MiniChat",
    "Hình 4. Mô phỏng giao diện sau khi kết nối và gửi tin nhắn",
]
TABLE_LABELS = [
    "Bảng 1. Các thành phần chính của ứng dụng MiniChat",
    "Bảng 2. Kết quả kiểm thử các chức năng chính",
]


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def write_list_entry(paragraph, label):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    run.font.size = Pt(10)
    page = paragraph.add_run("\t1")
    page.font.size = Pt(10)


def write_caption(paragraph, label):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(label)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)


def main():
    doc = Document(DOCX)
    seen = {label: 0 for label in FIGURE_LABELS + TABLE_LABELS}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for label in seen:
            if text.startswith(label):
                seen[label] += 1
                if seen[label] == 1:
                    write_list_entry(paragraph, label)
                else:
                    write_caption(paragraph, label)
                break

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
