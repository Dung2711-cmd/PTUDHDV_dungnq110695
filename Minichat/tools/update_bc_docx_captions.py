from pathlib import Path
from shutil import copy2
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "BC.docx"
BACKUP = ROOT / "BC_backup_before_caption_update.docx"


FIGURE_CAPTIONS = [
    "Hình 1. Kiến trúc client-server của MiniChat",
    "Hình 2. Mô hình giao tiếp client - server",
    "Hình 3. Luồng xử lý gửi và nhận tin nhắn trong MiniChat",
    "Hình 4. Mô phỏng giao diện sau khi kết nối và gửi tin nhắn",
]


def set_update_fields_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph, text, style_name=None, bold=False, italic=False, size=11, color=None):
    clear_paragraph(paragraph)
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            paragraph.style = "normal"
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def insert_paragraph_after(paragraph, text="", style_name=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def is_caption_candidate(text):
    normalized = text.strip().lower()
    normalized = normalized.replace("hình", "hình")
    return (
        normalized.startswith("hình")
        or normalized.startswith("kiến trúc client-server")
        or normalized.startswith("mô phỏng giao diện")
    )


def style_caption(paragraph, text):
    set_paragraph_text(paragraph, text, style_name="Caption", italic=True, size=10, color=(80, 80, 80))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)


def get_paragraphs(doc):
    return list(doc.paragraphs)


def update_figure_captions(doc):
    paragraphs = get_paragraphs(doc)
    image_indices = [i for i, p in enumerate(paragraphs) if p._p.xpath(".//w:drawing")]
    changed = []

    for number, idx in enumerate(image_indices[: len(FIGURE_CAPTIONS)]):
        paragraphs = get_paragraphs(doc)
        image_para = paragraphs[idx]
        caption_text = FIGURE_CAPTIONS[number]
        next_para = paragraphs[idx + 1] if idx + 1 < len(paragraphs) else None

        if next_para is not None and is_caption_candidate(next_para.text):
            style_caption(next_para, caption_text)
            changed.append(caption_text)
        else:
            new_caption = insert_paragraph_after(image_para)
            style_caption(new_caption, caption_text)
            changed.append(caption_text)

    return changed


def find_heading(doc, title):
    needle = title.strip().lower()
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().lower() == needle:
            return paragraph
    return None


def remove_section_body(doc, heading_title, stop_headings):
    paragraphs = get_paragraphs(doc)
    heading_index = None
    for i, p in enumerate(paragraphs):
        if p.text.strip().lower() == heading_title.lower():
            heading_index = i
            break
    if heading_index is None:
        return None

    stop_titles = {x.lower() for x in stop_headings}
    to_delete = []
    for p in paragraphs[heading_index + 1 :]:
        text = p.text.strip().lower()
        if text in stop_titles:
            break
        to_delete.append(p)

    for p in to_delete:
        delete_paragraph(p)

    return find_heading(doc, heading_title)


def heading_level(style_name):
    match = re.search(r"Heading ([1-3])", style_name)
    if not match:
        return None
    return int(match.group(1))


def collect_headings(doc):
    entries = []
    skip_titles = {"mục lục", "danh mục hình ảnh", "danh mục bảng"}
    for p in doc.paragraphs:
        text = p.text.strip()
        level = heading_level(p.style.name)
        if not text or level is None:
            continue
        if text.lower() in skip_titles:
            continue
        entries.append((level, text))
    return entries


def add_toc_field(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph._p.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    paragraph._p.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph._p.append(separate)

    run = paragraph.add_run("Mục lục sẽ được Word tự cập nhật khi mở tài liệu.")
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(end)


def update_table_of_contents(doc):
    heading = remove_section_body(doc, "MỤC LỤC", ["DANH MỤC HÌNH ẢNH", "DANH MỤC BẢNG", "TÓM TẮT"])
    if heading is None:
        return

    field_para = insert_paragraph_after(heading)
    add_toc_field(field_para)

    # Also provide a visible outline so the section is useful before Word refreshes field codes.
    last = field_para
    for level, text in collect_headings(doc):
        p = insert_paragraph_after(last)
        p.paragraph_format.left_indent = Pt((level - 1) * 18)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10)
        if level == 1:
            run.bold = True
        last = p


def update_list_of_figures(doc):
    heading = remove_section_body(doc, "DANH MỤC HÌNH ẢNH", ["DANH MỤC BẢNG", "TÓM TẮT"])
    if heading is None:
        return
    last = heading
    for caption in FIGURE_CAPTIONS:
        p = insert_paragraph_after(last)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(caption)
        last = p


def ensure_list_of_tables_heading(doc):
    table_heading = find_heading(doc, "DANH MỤC BẢNG")
    if table_heading is not None:
        return table_heading

    figure_heading = find_heading(doc, "DANH MỤC HÌNH ẢNH")
    if figure_heading is None:
        summary_heading = find_heading(doc, "TÓM TẮT")
        if summary_heading is None:
            return None
        new_heading = insert_paragraph_after(summary_heading, "DANH MỤC BẢNG", "Heading 1")
        return new_heading

    paragraphs = get_paragraphs(doc)
    idx = next(
        (
            i
            for i, p in enumerate(paragraphs)
            if p.text.strip().lower() == "danh mục hình ảnh"
        ),
        None,
    )
    if idx is None:
        return insert_paragraph_after(figure_heading, "DANH MỤC BẢNG", "Heading 1")
    next_major = None
    for p in paragraphs[idx + 1 :]:
        if p.style.name.startswith("Heading") and p.text.strip().upper() in {"TÓM TẮT", "MỞ ĐẦU"}:
            next_major = p
            break
    if next_major is None:
        return insert_paragraph_after(figure_heading, "DANH MỤC BẢNG", "Heading 1")

    new_p = OxmlElement("w:p")
    next_major._p.addprevious(new_p)
    new_heading = Paragraph(new_p, next_major._parent)
    new_heading.style = "Heading 1"
    new_heading.add_run("DANH MỤC BẢNG")
    return new_heading


def update_list_of_tables(doc):
    heading = ensure_list_of_tables_heading(doc)
    if heading is None:
        return
    remove_section_body(doc, "DANH MỤC BẢNG", ["TÓM TẮT", "MỞ ĐẦU"])
    heading = find_heading(doc, "DANH MỤC BẢNG")
    p = insert_paragraph_after(heading)
    run = p.add_run("Tài liệu hiện chưa có bảng biểu cần đánh số.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(90, 90, 90)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not BACKUP.exists():
        copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    set_update_fields_on_open(doc)
    captions = update_figure_captions(doc)
    update_list_of_tables(doc)
    update_list_of_figures(doc)
    update_table_of_contents(doc)
    doc.save(DOCX)

    print(f"Updated: {DOCX}")
    print(f"Backup: {BACKUP}")
    print(f"Figure captions: {len(captions)}")
    print(f"Tables in document: {len(doc.tables)}")


if __name__ == "__main__":
    main()
