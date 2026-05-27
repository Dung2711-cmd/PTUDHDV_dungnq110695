from pathlib import Path
from shutil import copy2
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "BC.docx"
BACKUP = ROOT / "BC_backup_before_table_list_update.docx"

FIGURE_CAPTIONS = [
    "Hình 1. Kiến trúc client-server của MiniChat",
    "Hình 2. Mô hình giao tiếp client - server",
    "Hình 3. Luồng xử lý gửi và nhận tin nhắn trong MiniChat",
    "Hình 4. Mô phỏng giao diện sau khi kết nối và gửi tin nhắn",
]

TABLES_TO_ADD = [
    {
        "anchor": "1.1 Các thành phần chính",
        "caption": "Bảng 1. Các thành phần chính của ứng dụng MiniChat",
        "bookmark": "tbl_components",
        "rows": [
            ["Thành phần", "Vai trò"],
            ["ChatServer.java", "Khởi tạo ServerSocket tại cổng 9999, tiếp nhận client và broadcast tin nhắn."],
            ["ClientHandler.java", "Xử lý một client trong luồng riêng: nhận username, đọc tin nhắn và đóng kết nối."],
            ["ChatClient.java", "Client dòng lệnh dùng để kiểm thử nhanh trong terminal."],
            ["ChatClientGUI.java", "Client giao diện Swing, hỗ trợ nhập username, kết nối và gửi tin nhắn."],
            ["module-info.java", "Khai báo module Minichat và quyền dùng java.desktop cho giao diện Swing."],
        ],
        "widths": [1.8, 4.4],
    },
    {
        "anchor": "CHƯƠNG 4: ĐÁNH GIÁ KẾT QUẢ",
        "caption": "Bảng 2. Kết quả kiểm thử các chức năng chính",
        "bookmark": "tbl_testing",
        "rows": [
            ["Mã", "Kịch bản kiểm thử", "Kết quả"],
            ["TC01", "Chạy ChatServer trước khi mở client.", "Đạt"],
            ["TC02", "Chạy client khi server chưa chạy.", "Client báo lỗi kết nối rõ ràng."],
            ["TC03", "Mở hai client GUI và gửi tin nhắn.", "Tin nhắn được chuyển đến client còn lại."],
            ["TC04", "Người dùng nhập /quit hoặc bấm Disconnect.", "Client ngắt kết nối và server xóa khỏi danh sách."],
            ["TC05", "Chạy nhiều cửa sổ client trên cùng máy.", "Server tạo ClientHandler riêng cho từng client."],
        ],
        "widths": [0.8, 3.6, 1.8],
    },
]


def paragraphs(doc):
    return list(doc.paragraphs)


def set_update_fields_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def insert_paragraph_after(paragraph, text="", style_name=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except KeyError:
            new_para.style = "normal"
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc, text):
    target = text.strip().lower()
    for p in doc.paragraphs:
        if p.text.strip().lower() == target:
            return p
    return None


def find_last_paragraph(doc, text):
    target = text.strip().lower()
    found = None
    for p in doc.paragraphs:
        if p.text.strip().lower() == target:
            found = p
    return found


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def insert_table_after(doc, anchor_paragraph, rows, widths):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = False
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)

    for cell, value in zip(table.rows[0].cells, rows[0]):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, value, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor_paragraph._p.addnext(tbl)
    return table


def add_bookmark(paragraph, name, bookmark_id):
    existing = paragraph._p.xpath(f'.//w:bookmarkStart[@w:name="{name}"]')
    if existing:
        return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def style_caption(paragraph, text, bookmark, bookmark_id):
    clear_paragraph(paragraph)
    paragraph.style = "normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    add_bookmark(paragraph, bookmark, bookmark_id)


def ensure_tables(doc):
    added = []
    for number, spec in enumerate(TABLES_TO_ADD, start=1):
        if find_paragraph(doc, spec["caption"]):
            continue
        anchor = find_last_paragraph(doc, spec["anchor"])
        if anchor is None:
            continue
        caption_para = insert_paragraph_after(anchor)
        style_caption(caption_para, spec["caption"], spec["bookmark"], 200 + number)
        insert_table_after(doc, caption_para, spec["rows"], spec["widths"])
        added.append(spec["caption"])
    return added


def remove_all_tables(doc):
    for table in list(doc.tables):
        tbl = table._tbl
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)


def remove_orphan_table_caption_blanks(doc):
    for p in list(doc.paragraphs):
        if p.text.strip() in {spec["caption"] for spec in TABLES_TO_ADD}:
            delete_paragraph(p)


def normalize_figure_captions(doc):
    image_paras = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
    for idx, image_para in enumerate(image_paras[: len(FIGURE_CAPTIONS)], start=1):
        all_paras = paragraphs(doc)
        pos = next((i for i, p in enumerate(all_paras) if p._p is image_para._p), None)
        if pos is None:
            continue
        next_para = all_paras[pos + 1] if pos + 1 < len(all_paras) else None
        caption = FIGURE_CAPTIONS[idx - 1]
        if next_para is not None and (
            next_para.text.strip().startswith("Hình")
            or next_para.text.strip().startswith("Kiến trúc client-server")
            or next_para.text.strip().startswith("Mô phỏng giao diện")
        ):
            style_caption(next_para, caption, f"fig_{idx}", 100 + idx)
        else:
            p = insert_paragraph_after(image_para)
            style_caption(p, caption, f"fig_{idx}", 100 + idx)


def remove_section_body(doc, heading_title, stop_headings):
    all_paras = paragraphs(doc)
    heading_index = next((i for i, p in enumerate(all_paras) if p.text.strip().lower() == heading_title.lower()), None)
    if heading_index is None:
        return None
    stops = {s.lower() for s in stop_headings}
    to_delete = []
    for p in all_paras[heading_index + 1 :]:
        if p.text.strip().lower() in stops:
            break
        to_delete.append(p)
    for p in to_delete:
        delete_paragraph(p)
    return find_paragraph(doc, heading_title)


def ensure_heading_before(doc, heading_text, before_text):
    heading = find_paragraph(doc, heading_text)
    if heading:
        return heading
    before = find_paragraph(doc, before_text)
    if before is None:
        return None
    new_p = OxmlElement("w:p")
    before._p.addprevious(new_p)
    heading = Paragraph(new_p, before._parent)
    heading.style = "Heading 1"
    heading.add_run(heading_text)
    return heading


def add_pageref_field(paragraph, bookmark):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph._p.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    paragraph._p.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph._p.append(separate)

    run = paragraph.add_run("1")
    run.font.size = Pt(10)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(end)


def add_list_entry(after_paragraph, label, bookmark):
    p = insert_paragraph_after(after_paragraph)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.size = Pt(10)
    p.add_run("\t")
    add_pageref_field(p, bookmark)
    return p


def update_list_of_figures(doc):
    heading = remove_section_body(doc, "DANH MỤC HÌNH ẢNH", ["DANH MỤC BẢNG", "TÓM TẮT"])
    if heading is None:
        return
    last = heading
    for idx, caption in enumerate(FIGURE_CAPTIONS, start=1):
        last = add_list_entry(last, caption, f"fig_{idx}")


def update_list_of_tables(doc):
    heading = ensure_heading_before(doc, "DANH MỤC BẢNG", "TÓM TẮT")
    if heading is None:
        return
    remove_section_body(doc, "DANH MỤC BẢNG", ["TÓM TẮT"])
    heading = find_paragraph(doc, "DANH MỤC BẢNG")
    last = heading
    for spec in TABLES_TO_ADD:
        last = add_list_entry(last, spec["caption"], spec["bookmark"])


def add_toc_field(paragraph):
    clear_paragraph(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph._p.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    paragraph._p.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph._p.append(separate)
    r = paragraph.add_run("Mục lục sẽ được Word tự cập nhật khi mở tài liệu.")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(end)


def heading_level(style_name):
    m = re.search(r"Heading ([1-3])", style_name)
    return int(m.group(1)) if m else None


def update_toc(doc):
    heading = remove_section_body(doc, "MỤC LỤC", ["DANH MỤC HÌNH ẢNH"])
    if heading is None:
        return
    field_para = insert_paragraph_after(heading)
    add_toc_field(field_para)
    last = field_para
    skip = {"mục lục"}
    for p in doc.paragraphs:
        text = p.text.strip()
        level = heading_level(p.style.name)
        if not text or level is None or text.lower() in skip:
            continue
        entry = insert_paragraph_after(last)
        entry.paragraph_format.left_indent = Pt((level - 1) * 18)
        entry.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = entry.add_run(text)
        run.font.size = Pt(10)
        if level == 1:
            run.bold = True
        entry.add_run("\t")
        # Mục lục thật được Word cập nhật bằng field ở trên; dòng dưới là khung xem nhanh.
        pr = entry.add_run("...")
        pr.font.size = Pt(10)
        last = entry


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not BACKUP.exists():
        copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    set_update_fields_on_open(doc)
    remove_all_tables(doc)
    remove_orphan_table_caption_blanks(doc)
    normalize_figure_captions(doc)
    update_list_of_figures(doc)
    update_list_of_tables(doc)
    update_toc(doc)
    added = ensure_tables(doc)
    doc.save(DOCX)

    print(f"Updated: {DOCX}")
    print(f"Backup: {BACKUP}")
    print(f"Added tables: {len(added)}")
    print(f"Tables now: {len(Document(DOCX).tables)}")


if __name__ == "__main__":
    main()
