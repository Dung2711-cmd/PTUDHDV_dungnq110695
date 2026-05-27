from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_assets"
OUT_DOCX = ROOT / "Bao_cao_de_tai_MiniChat.docx"


INK = (30, 41, 59)
MUTED = (100, 116, 139)
BLUE = (37, 99, 235)
TEAL = (13, 148, 136)
GREEN = (22, 163, 74)
ORANGE = (234, 88, 12)
PANEL = (248, 250, 252)
BORDER = (203, 213, 225)


def font(size=24, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, xy, radius=18, fill=(255, 255, 255), outline=BORDER, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, fill=INK, size=24, bold=False):
    fnt = font(size, bold)
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap(raw, 24) or [""])
    line_h = int(size * 1.25)
    total_h = line_h * len(lines)
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = box[0] + ((box[2] - box[0]) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h


def arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        head = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
    else:
        head = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    draw.polygon(head, fill=color)


def draw_architecture(path):
    img = Image.new("RGB", (1400, 760), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1400, 760), fill=(246, 248, 251))
    d.text((70, 50), "Kiến trúc tổng quan MiniChat", font=font(42, True), fill=INK)
    d.text((70, 105), "Mô hình client-server: nhiều client kết nối đến một server qua TCP socket.", font=font(24), fill=MUTED)

    client1 = (90, 250, 385, 440)
    server = (555, 220, 845, 485)
    client2 = (1015, 250, 1310, 440)
    for box, title, detail, color in [
        (client1, "Client A", "ChatClientGUI\nSocket localhost:9999", BLUE),
        (server, "ChatServer", "ServerSocket 9999\nTạo ClientHandler", TEAL),
        (client2, "Client B", "ChatClientGUI\nGửi/nhận tin nhắn", GREEN),
    ]:
        rounded(d, box, radius=26, fill=(255, 255, 255), outline=color, width=4)
        d.text((box[0] + 30, box[1] + 28), title, font=font(32, True), fill=color)
        centered_text(d, (box[0] + 20, box[1] + 75, box[2] - 20, box[3] - 20), detail, fill=INK, size=24)

    arrow(d, (385, 330), (555, 330), BLUE)
    arrow(d, (845, 330), (1015, 330), GREEN)
    arrow(d, (555, 400), (385, 400), BLUE)
    arrow(d, (1015, 400), (845, 400), GREEN)
    d.text((470, 292), "TCP", font=font(22, True), fill=BLUE)
    d.text((910, 292), "TCP", font=font(22, True), fill=GREEN)
    d.text((472, 430), "broadcast", font=font(22), fill=MUTED)
    d.text((882, 430), "broadcast", font=font(22), fill=MUTED)

    rounded(d, (260, 590, 1140, 675), radius=18, fill=(239, 246, 255), outline=(147, 197, 253), width=2)
    centered_text(d, (290, 595, 1110, 670), "Server giữ danh sách client và chuyển tiếp tin nhắn đến các client còn lại.", fill=INK, size=26)
    img.save(path)


def draw_flow(path):
    img = Image.new("RGB", (1400, 840), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1400, 840), fill=(248, 250, 252))
    d.text((70, 50), "Luồng xử lý khi gửi tin nhắn", font=font(42, True), fill=INK)
    steps = [
        ("1", "Người dùng nhập username và bấm Connect", BLUE),
        ("2", "Client mở Socket đến localhost:9999", TEAL),
        ("3", "Server accept kết nối và tạo ClientHandler", GREEN),
        ("4", "Client gửi nội dung tin nhắn", ORANGE),
        ("5", "Server broadcast đến các client khác", BLUE),
        ("6", "Client nhận hiển thị lên vùng chat", TEAL),
    ]
    x, y = 105, 150
    w, h = 540, 120
    coords = []
    for i, (num, text, color) in enumerate(steps):
        col = i % 2
        row = i // 2
        bx = x + col * 650
        by = y + row * 190
        coords.append((bx, by, bx + w, by + h))
        rounded(d, coords[-1], radius=22, fill="white", outline=color, width=4)
        d.ellipse((bx + 28, by + 30, bx + 88, by + 90), fill=color)
        centered_text(d, (bx + 28, by + 30, bx + 88, by + 90), num, fill="white", size=28, bold=True)
        d.text((bx + 112, by + 36), text, font=font(25, True), fill=INK)
    for i in range(len(coords) - 1):
        a = coords[i]
        b = coords[i + 1]
        start = (a[2] - 8, (a[1] + a[3]) // 2) if i % 2 == 0 else ((a[0] + a[2]) // 2, a[3] + 8)
        end = (b[0] + 8, (b[1] + b[3]) // 2) if i % 2 == 0 else ((b[0] + b[2]) // 2, b[1] - 8)
        arrow(d, start, end, MUTED, 4)
    img.save(path)


def draw_gui(path):
    img = Image.new("RGB", (1400, 840), (235, 239, 245))
    d = ImageDraw.Draw(img)
    rounded(d, (135, 85, 1265, 755), radius=34, fill=(248, 250, 252), outline=(190, 196, 205), width=3)
    d.rectangle((136, 145, 1264, 754), fill=(248, 250, 252))
    for x, color in [(175, (255, 95, 87)), (225, (255, 189, 46)), (275, (40, 201, 64))]:
        d.ellipse((x, 112, x + 28, 140), fill=color, outline=(120, 120, 120))
    d.text((325, 106), "MiniChat", font=font(32, True), fill=(82, 82, 82))
    d.text((185, 190), "Username", font=font(28, True), fill=INK)
    d.rectangle((340, 180, 1005, 225), fill="white", outline=(226, 232, 240))
    d.text((355, 188), "An", font=font(25), fill=(100, 116, 139))
    rounded(d, (1040, 180, 1240, 225), radius=12, fill="white", outline=(203, 213, 225), width=2)
    centered_text(d, (1040, 180, 1240, 225), "Disconnect", fill=INK, size=25, bold=True)
    d.rectangle((185, 260, 1240, 620), fill="white", outline=(203, 213, 225), width=2)
    d.text((215, 292), "Binh: Xin chào An", font=font(28), fill=INK)
    d.text((215, 338), "Me: Chào Binh, mình đang test MiniChat", font=font(28), fill=INK)
    d.text((185, 650), "Connected to localhost:9999", font=font(26, True), fill=INK)
    d.rectangle((185, 695, 1075, 740), fill="white", outline=(100, 116, 139), width=2)
    d.text((198, 703), "Nhập tin nhắn...", font=font(25), fill=(100, 116, 139))
    rounded(d, (1115, 695, 1240, 740), radius=12, fill="white", outline=(203, 213, 225), width=2)
    centered_text(d, (1115, 695, 1240, 740), "Send", fill=INK, size=26, bold=True)
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=1, cols=len(widths))
    table.style = "Table Grid"
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    header = table.rows[0].cells
    for cell, text in zip(header, rows[0]):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    return table


def build_docx():
    OUT_DIR.mkdir(exist_ok=True)
    arch = OUT_DIR / "kien_truc_minichat.png"
    flow = OUT_DIR / "luong_xu_ly_tin_nhan.png"
    gui = OUT_DIR / "giao_dien_minichat.png"
    draw_architecture(arch)
    draw_flow(flow)
    draw_gui(gui)

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BÁO CÁO ĐỀ TÀI\nỨNG DỤNG CHAT MINI BẰNG JAVA SOCKET")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Đề tài: MiniChat - mô hình client-server, đa luồng và giao diện Java Swing")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    doc.add_picture(str(gui), width=Inches(5.9))
    add_caption(doc, "Hình 1. Giao diện minh họa của ứng dụng MiniChat")
    doc.add_page_break()

    doc.add_heading("1. Tóm tắt", level=1)
    doc.add_paragraph(
        "Đề tài MiniChat xây dựng một ứng dụng trò chuyện đơn giản bằng Java, sử dụng mô hình client-server. "
        "Server lắng nghe kết nối tại cổng 9999, mỗi client sau khi kết nối sẽ được phục vụ bởi một luồng riêng thông qua lớp ClientHandler. "
        "Người dùng có thể nhập tên, gửi tin nhắn và nhận tin nhắn từ những người dùng khác. Ngoài client chạy trên terminal, đề tài đã bổ sung client giao diện bằng Java Swing để thao tác trực quan hơn."
    )

    doc.add_heading("2. Mở đầu", level=1)
    doc.add_paragraph(
        "Trong các hệ thống mạng, ứng dụng chat là ví dụ cơ bản nhưng rất hữu ích để tìm hiểu cách nhiều chương trình giao tiếp với nhau qua socket. "
        "Thông qua đề tài này, người thực hiện có thể nắm được nguyên lý tạo server, kết nối client, truyền dữ liệu văn bản, xử lý nhiều người dùng đồng thời và xây dựng giao diện người dùng cơ bản."
    )
    doc.add_heading("2.1. Mục tiêu đề tài", level=2)
    for item in [
        "Xây dựng server có khả năng tiếp nhận nhiều client cùng lúc.",
        "Xây dựng client gửi và nhận tin nhắn theo thời gian gần thực.",
        "Áp dụng lập trình đa luồng để mỗi client có một tiến trình xử lý độc lập.",
        "Thiết kế giao diện MiniChat bằng Java Swing, dễ thao tác hơn so với terminal.",
        "Kiểm tra tình huống lỗi thường gặp như server chưa chạy hoặc client không kết nối được.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Tổng quan", level=1)
    doc.add_paragraph(
        "MiniChat được tổ chức theo package minichat. Các thành phần chính gồm ChatServer, ClientHandler, ChatClient và ChatClientGUI. "
        "Trong đó ChatServer chịu trách nhiệm mở cổng kết nối, ClientHandler xử lý từng client, còn ChatClientGUI là giao diện người dùng."
    )
    doc.add_picture(str(arch), width=Inches(6.3))
    add_caption(doc, "Hình 2. Kiến trúc client-server của MiniChat")
    add_table(
        doc,
        [
            ("Thành phần", "Vai trò"),
            ("ChatServer.java", "Khởi tạo ServerSocket tại cổng 9999, chấp nhận client mới và broadcast tin nhắn."),
            ("ClientHandler.java", "Xử lý một client: nhận username, đọc tin nhắn, gửi tin nhắn đến client khác và đóng kết nối."),
            ("ChatClient.java", "Client dòng lệnh, phù hợp kiểm thử nhanh trong terminal."),
            ("ChatClientGUI.java", "Client giao diện Swing, có ô nhập tên, vùng chat, trạng thái kết nối và nút gửi."),
            ("module-info.java", "Khai báo module Minichat và quyền dùng java.desktop cho Swing."),
        ],
        [1.8, 4.4],
    )

    doc.add_heading("4. Nghiên cứu lý thuyết", level=1)
    doc.add_heading("4.1. Mô hình client-server", level=2)
    doc.add_paragraph(
        "Mô hình client-server gồm một chương trình server cung cấp dịch vụ và nhiều chương trình client kết nối đến để sử dụng dịch vụ đó. "
        "Trong MiniChat, server đóng vai trò trung tâm chuyển tiếp tin nhắn giữa các client."
    )
    doc.add_heading("4.2. TCP Socket trong Java", level=2)
    doc.add_paragraph(
        "Java cung cấp ServerSocket để server lắng nghe kết nối và Socket để client/server trao đổi dữ liệu. "
        "Dữ liệu trong đề tài được truyền dưới dạng văn bản qua BufferedReader và PrintWriter."
    )
    doc.add_heading("4.3. Đa luồng", level=2)
    doc.add_paragraph(
        "Nếu server chỉ chạy một luồng, nó khó phục vụ nhiều client cùng lúc. Vì vậy, mỗi khi có client kết nối, ChatServer tạo một ClientHandler và chạy trong một Thread riêng. "
        "Cách này giúp nhiều người dùng có thể gửi tin nhắn đồng thời."
    )
    doc.add_heading("4.4. Java Swing", level=2)
    doc.add_paragraph(
        "Swing là thư viện giao diện đi kèm Java. Trong đề tài, Swing được dùng để tạo cửa sổ MiniChat gồm trường username, vùng hiển thị tin nhắn, ô nhập nội dung, nút Connect/Disconnect và nút Send."
    )

    doc.add_heading("5. Nội dung thực hiện", level=1)
    doc.add_heading("5.1. Xây dựng server", level=2)
    doc.add_paragraph(
        "Server sử dụng ServerSocket tại cổng 9999. Khi có client kết nối, server gọi accept(), tạo đối tượng ClientHandler, thêm vào danh sách clients và khởi động một Thread mới để xử lý client đó."
    )
    doc.add_heading("5.2. Xử lý client", level=2)
    doc.add_paragraph(
        "ClientHandler gửi yêu cầu nhập username, sau đó liên tục đọc tin nhắn từ client. Khi nhận tin nhắn hợp lệ, server gọi hàm broadcast để gửi nội dung đó đến những client khác. Nếu client nhập /quit hoặc mất kết nối, server xóa client khỏi danh sách."
    )
    doc.add_heading("5.3. Xây dựng giao diện", level=2)
    doc.add_paragraph(
        "Giao diện ChatClientGUI được xây dựng bằng JFrame, JTextArea, JTextField, JButton và JLabel. "
        "Khi người dùng bấm Connect, chương trình mở socket đến localhost:9999, gửi username và tạo luồng nhận tin nhắn từ server. "
        "Các thông báo kỹ thuật như 'Enter your username:' được lọc khỏi vùng chat để giao diện sạch hơn."
    )
    doc.add_picture(str(flow), width=Inches(6.3))
    add_caption(doc, "Hình 3. Luồng xử lý gửi và nhận tin nhắn")
    doc.add_heading("5.4. Các bước chạy chương trình", level=2)
    for step in [
        "Biên dịch project: javac -d bin src/module-info.java src/minichat/*.java",
        "Chạy server trước: java -p bin -m Minichat/minichat.ChatServer",
        "Chạy client giao diện: java -p bin -m Minichat/minichat.ChatClientGUI",
        "Nhập username, bấm Connect, sau đó nhập tin nhắn và bấm Send.",
        "Mở thêm client thứ hai để kiểm thử việc trao đổi tin nhắn giữa nhiều người dùng.",
    ]:
        add_number(doc, step)

    doc.add_heading("6. Đánh giá kết quả", level=1)
    doc.add_paragraph(
        "Ứng dụng đã đáp ứng các chức năng cơ bản của một hệ thống chat nội bộ. Server có thể nhận nhiều client, phân biệt người gửi bằng username và chuyển tiếp tin nhắn đến các client còn lại. "
        "Giao diện Swing giúp thao tác thuận tiện hơn so với client dòng lệnh."
    )
    add_table(
        doc,
        [
            ("Tiêu chí", "Kết quả đánh giá"),
            ("Kết nối client-server", "Hoạt động khi server được chạy trước tại localhost:9999."),
            ("Gửi và nhận tin nhắn", "Client gửi tin nhắn lên server; server broadcast đến client khác."),
            ("Đa người dùng", "Mỗi client được xử lý bởi một Thread riêng."),
            ("Giao diện", "Có cửa sổ chat, trạng thái kết nối, ô nhập username và nút gửi."),
            ("Xử lý lỗi", "Client báo rõ khi server chưa chạy hoặc không kết nối được."),
            ("Hạn chế", "Chưa có đăng nhập thật, lưu lịch sử chat, mã hóa dữ liệu hoặc chat nhóm hoàn chỉnh."),
        ],
        [2.0, 4.2],
    )
    doc.add_heading("6.1. Hướng phát triển", level=2)
    for item in [
        "Bổ sung danh sách người dùng đang online.",
        "Thêm chức năng chat riêng và chat nhóm.",
        "Lưu lịch sử tin nhắn vào file hoặc cơ sở dữ liệu.",
        "Cải thiện giao diện theo hướng hiện đại hơn.",
        "Bổ sung mã hóa hoặc xác thực để tăng tính an toàn.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7. Tài liệu tham khảo", level=1)
    references = [
        "Oracle. Java Platform, Standard Edition Documentation - java.net.ServerSocket và java.net.Socket.",
        "Oracle. Java Swing Tutorial - Creating a GUI with Swing.",
        "Oracle. Java I/O Documentation - BufferedReader, InputStreamReader và PrintWriter.",
        "Tài liệu bài học Lập trình mạng: mô hình client-server và lập trình socket TCP.",
        "Mã nguồn project MiniChat trong thư mục src/minichat.",
    ]
    for ref in references:
        add_number(doc, ref)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Báo cáo đề tài MiniChat")

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_docx())
